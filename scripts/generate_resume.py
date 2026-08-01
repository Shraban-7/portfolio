"""
Generate ATS-friendly resume (DOCX + PDF) aligned with portfolio experience.
Run: python scripts/generate_resume.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]

NAME = "Shraban"
ROLE = "Full Stack Developer"
EMAIL = "shakuatshraban@gmail.com"
PHONE = ""  # add if desired
LOCATION = "Bangladesh"
GITHUB = "https://github.com/Shraban-7"
WEBSITE = "shraban.dev"

SUMMARY = (
    "Full Stack Developer with hands-on experience building production web applications "
    "using Laravel, PHP, JavaScript/TypeScript, and modern frontends (Next.js, Vue, Livewire). "
    "Delivered E-commerce platforms, POS systems, SaaS inventory tools, HRM, school management, "
    "LMS, real-estate, and CMS products across agency, freelance, and product-company roles. "
    "Focused on reliable backends, clean APIs, security, and maintainable full-stack delivery."
)

SKILLS = {
    "Languages": "PHP, JavaScript, TypeScript",
    "Frameworks": "Laravel, Next.js, Vue 3, Inertia.js, Livewire, Alpine.js, Tailwind CSS, Blade",
    "Databases": "MySQL, PostgreSQL, SQLite, Redis",
    "Tools": "Git, GitHub, Docker, Vite, Laravel Sanctum, Pest, PHPUnit, REST APIs",
    "Domains": "E-commerce, POS, SaaS, HRM, LMS, School Management, Real Estate, CMS",
}

EXPERIENCE = [
    {
        "title": "Full Stack Developer",
        "org": "Spinner Tech",
        "period": "January 2025 – Present",
        "bullets": [
            "Develop and maintain E-commerce platforms, POS systems, and shopping/task-related websites for production use.",
            "Build real-estate management applications and social-media marketing CMS features used by business clients.",
            "Contribute to fantasy game products and related web features with Laravel backends and modern UI stacks.",
            "Collaborate on full-stack delivery across PHP, Laravel, JavaScript/TypeScript, and responsive frontends.",
        ],
    },
    {
        "title": "Freelance Developer",
        "org": "Freelance",
        "period": "February 2024 – December 2024",
        "bullets": [
            "Delivered university and personal software projects with a focus on practical, secure application design.",
            "Built a security guard management application for operational tracking and administration workflows.",
            "Implemented an NID encryption project emphasizing secure handling of sensitive identity data.",
        ],
    },
    {
        "title": "Full Stack Developer",
        "org": "Web Arts Factory",
        "period": "March 2023 – February 2024",
        "bullets": [
            "Delivered Point of Sale (POS) systems for retail and business operations.",
            "Built agency websites optimized for presentation, conversion, and client branding needs.",
            "Developed a bank document management application for structured document workflows.",
        ],
    },
]

EDUCATION = {
    "degree": "Bachelor of Science in Computer Science & Engineering",
    "school": "Mymensingh Engineering College",
    "period": "2019 – 2024",
    "details": "Coursework in software engineering, databases, AI fundamentals, and web application development.",
}

PROJECTS = [
    {
        "name": "Multi-Vendor E-Commerce",
        "stack": "Laravel 11, PHP, Blade, Tailwind CSS, Sanctum, MySQL",
        "blurb": "Marketplace backend with seller portals, commissions, variants, bulk import, and multi-warehouse shipping.",
        "url": "https://github.com/Shraban-7/multivendor-ecommerce",
    },
    {
        "name": "Restaurant POS",
        "stack": "Laravel 11, Alpine.js, Tailwind CSS, Reverb, MySQL, Docker",
        "blurb": "Restaurant POS with kitchen display, QR ordering, recipe BOM stock deduction, and offline PWA sync.",
        "url": "https://github.com/Shraban-7/resturant-pos",
    },
    {
        "name": "Inventory SaaS",
        "stack": "Laravel 13, Next.js, TypeScript, Tailwind CSS, Spatie Permission, MySQL",
        "blurb": "Multi-tenant inventory ERP with sales/purchasing, double-entry ledger, and RBAC.",
        "url": "https://github.com/Shraban-7/inventory-saas",
    },
    {
        "name": "Enterprise HRM",
        "stack": "Laravel 13, Blade, Tailwind CSS, jQuery, Vite, MySQL",
        "blurb": "HR platform for payroll/tax, attendance/leave, ATS careers portal, reviews, and audit trails.",
        "url": "https://github.com/Shraban-7/HRM",
    },
    {
        "name": "School Management System",
        "stack": "Laravel 13, Inertia.js, Vue 3, Tailwind CSS, MySQL",
        "blurb": "Single-school SMS with public site, academics, exams, fees, attendance, and bilingual UI.",
        "url": "https://github.com/Shraban-7/school-management-system",
    },
]


def _set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def build_docx(path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # Header — plain text for ATS
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    r = name_p.add_run(NAME.upper())
    _set_run_font(r, size=18, bold=True)

    role_p = doc.add_paragraph()
    role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_p.paragraph_format.space_after = Pt(2)
    r = role_p.add_run(ROLE)
    _set_run_font(r, size=12, bold=True)

    contact_bits = [EMAIL, LOCATION, GITHUB, WEBSITE]
    if PHONE:
        contact_bits.insert(1, PHONE)
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(10)
    r = contact_p.add_run(" | ".join(contact_bits))
    _set_run_font(r, size=10)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text.upper())
        _set_run_font(r, size=12, bold=True)
        # Underline via bottom border-like spacing: simple underline character line
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run("_" * 78)
        _set_run_font(r2, size=8, color=RGBColor(0x66, 0x66, 0x66))

    heading("Professional Summary")
    p = doc.add_paragraph(SUMMARY)
    for run in p.runs:
        _set_run_font(run, size=11)

    heading("Skills")
    for label, value in SKILLS.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        _set_run_font(r, size=11, bold=True)
        r = p.add_run(value)
        _set_run_font(r, size=11)

    heading("Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{job['title']} | {job['org']}")
        _set_run_font(r, size=11, bold=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(job["period"])
        _set_run_font(r, size=10)

        for bullet in job["bullets"]:
            bp = doc.add_paragraph(bullet, style="List Bullet")
            bp.paragraph_format.space_after = Pt(1)
            for run in bp.runs:
                _set_run_font(run, size=11)

    heading("Projects")
    for proj in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(proj["name"])
        _set_run_font(r, size=11, bold=True)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"Tech: {proj['stack']}")
        _set_run_font(r, size=10)

        p = doc.add_paragraph(proj["blurb"])
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            _set_run_font(run, size=11)

        p = doc.add_paragraph(proj["url"])
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            _set_run_font(run, size=9)

    heading("Education")
    p = doc.add_paragraph()
    r = p.add_run(f"{EDUCATION['degree']} | {EDUCATION['school']}")
    _set_run_font(r, size=11, bold=True)
    p = doc.add_paragraph(EDUCATION["period"])
    for run in p.runs:
        _set_run_font(run, size=10)
    p = doc.add_paragraph(EDUCATION["details"])
    for run in p.runs:
        _set_run_font(run, size=11)

    doc.save(path)


def build_pdf(path: Path) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    styles = getSampleStyleSheet()

    styles.add(
        ParagraphStyle(
            name="Name",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            alignment=1,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Role",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            alignment=1,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Contact",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            alignment=1,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11,
            spaceBefore=10,
            spaceAfter=4,
            textColor="#111111",
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="JobTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            spaceBefore=6,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="JobMeta",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBullet",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=12,
            leftIndent=12,
            spaceAfter=1,
        )
    )

    story = []
    story.append(Paragraph(NAME.upper(), styles["Name"]))
    story.append(Paragraph(ROLE, styles["Role"]))
    contact_bits = [EMAIL, LOCATION, GITHUB, WEBSITE]
    if PHONE:
        contact_bits.insert(1, PHONE)
    story.append(Paragraph(" | ".join(contact_bits), styles["Contact"]))

    story.append(Paragraph("PROFESSIONAL SUMMARY", styles["H"]))
    story.append(Paragraph("_" * 92, styles["JobMeta"]))
    story.append(Paragraph(SUMMARY, styles["Body"]))

    story.append(Paragraph("SKILLS", styles["H"]))
    story.append(Paragraph("_" * 92, styles["JobMeta"]))
    for label, value in SKILLS.items():
        story.append(Paragraph(f"<b>{label}:</b> {value}", styles["Body"]))

    story.append(Paragraph("EXPERIENCE", styles["H"]))
    story.append(Paragraph("_" * 92, styles["JobMeta"]))
    for job in EXPERIENCE:
        story.append(Paragraph(f"{job['title']} | {job['org']}", styles["JobTitle"]))
        story.append(Paragraph(job["period"], styles["JobMeta"]))
        for bullet in job["bullets"]:
            story.append(Paragraph(f"• {bullet}", styles["ResumeBullet"]))

    story.append(Paragraph("PROJECTS", styles["H"]))
    story.append(Paragraph("_" * 92, styles["JobMeta"]))
    for proj in PROJECTS:
        story.append(Paragraph(proj["name"], styles["JobTitle"]))
        story.append(Paragraph(f"Tech: {proj['stack']}", styles["JobMeta"]))
        story.append(Paragraph(proj["blurb"], styles["Body"]))
        story.append(Paragraph(proj["url"], styles["JobMeta"]))

    story.append(Paragraph("EDUCATION", styles["H"]))
    story.append(Paragraph("_" * 92, styles["JobMeta"]))
    story.append(
        Paragraph(
            f"{EDUCATION['degree']} | {EDUCATION['school']}",
            styles["JobTitle"],
        )
    )
    story.append(Paragraph(EDUCATION["period"], styles["JobMeta"]))
    story.append(Paragraph(EDUCATION["details"], styles["Body"]))

    story.append(Spacer(1, 8))
    doc.build(story)


def main() -> None:
    docx_path = ROOT / "Shraban_Full_Stack_Developer_Resume.docx"
    pdf_path = ROOT / "Shraban_Full_Stack_Developer_Resume.pdf"
    public_pdf = ROOT / "public" / "resume.pdf"

    build_docx(docx_path)
    build_pdf(pdf_path)
    build_pdf(public_pdf)

    print(f"Wrote {docx_path}")
    print(f"Wrote {pdf_path}")
    print(f"Wrote {public_pdf} (portfolio Resume button)")


if __name__ == "__main__":
    main()
